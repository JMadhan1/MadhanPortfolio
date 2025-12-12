from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_heading('J MADHAN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(24)
title_run.font.color.rgb = RGBColor(99, 102, 241)

# Subtitle
subtitle = doc.add_paragraph('AI/ML Developer | Full-Stack Developer')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(12)
subtitle_run.bold = True

# Contact Info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_text = '📧 jmadhan1@gmail.com | 🔗 GitHub: github.com/JMadhan1 | 💼 LinkedIn: linkedin.com/in/jmadhan1'
contact_run = contact.add_run(contact_text)
contact_run.font.size = Pt(9)

doc.add_paragraph('_' * 100)

# Professional Summary
doc.add_heading('PROFESSIONAL SUMMARY', 1)
summary = doc.add_paragraph(
    'Passionate AI/ML Developer with proven expertise in building production-ready solutions that solve real-world problems. '
    'Winner of multiple prestigious hackathons including IIIT Bangalore and MIT Mysore, with hands-on experience in computer vision, NLP, and cloud deployment. '
    'Delivered measurable impact through innovative projects achieving 99.4% accuracy in facial authenticity verification and 95% OCR accuracy in receipt processing systems.'
)
summary.runs[0].font.size = Pt(10)

highlights = doc.add_paragraph()
highlights_run = highlights.add_run('Key Highlights: ')
highlights_run.bold = True
highlights_run.font.size = Pt(10)
highlights_text = highlights.add_run('3x Hackathon Winner | 99.4% Model Accuracy | Team Leadership | Cloud Deployment Expert')
highlights_text.font.size = Pt(10)

# Technical Skills
doc.add_heading('TECHNICAL SKILLS', 1)

skills_data = [
    ('Expert Level:', 'Python • TensorFlow • Computer Vision • Machine Learning • Flask • Google Cloud'),
    ('Proficient:', 'PyTorch • Scikit-learn • OpenCV • NLP • Azure • FastAPI • Git'),
    ('Working Knowledge:', 'Pandas • NumPy • Docker • SQL • HTML5/CSS3 • JavaScript'),
    ('Learning:', 'Generative AI • Web3.py • MCP Protocol • Solana SDK • BERT • Prompt Engineering')
]

for label, skills in skills_data:
    p = doc.add_paragraph()
    p_label = p.add_run(label)
    p_label.bold = True
    p_label.font.size = Pt(10)
    p_skills = p.add_run(' ' + skills)
    p_skills.font.size = Pt(9)

# Featured Projects
doc.add_heading('FEATURED PROJECTS', 1)

projects = [
    {
        'name': '1. CodeMind AI - Intelligent Code Analysis Platform',
        'badge': '🏆 Featured | Live',
        'desc': 'AI-powered code analysis platform with real-time code reviews, smart suggestions, and automated refactoring',
        'tech': 'Python, AI/ML, Monaco Editor, Flask, GitHub API',
        'live': 'https://vibe2-1.onrender.com/',
        'github': 'https://github.com/JMadhan1/CodeMind'
    },
    {
        'name': '2. CiteCraft AI - Citation Management System',
        'badge': '🏆 MIT Mysore Winner | Live',
        'desc': 'Revolutionary AI-driven citation management system with automatic generation in multiple formats (APA, MLA, Chicago, Harvard)',
        'tech': 'HTML5, CSS3, Flask, FastAPI, PostgreSQL',
        'live': 'https://citecraft-ai.onrender.com/',
        'github': 'https://github.com/JMadhan1/CiteCraft.AI'
    },
    {
        'name': '3. InvestGuard AI - Investment Portfolio Management',
        'badge': '🏆 Featured | Live',
        'desc': 'Comprehensive AI-powered investment platform with real-time market analysis, risk assessment, and predictive analytics',
        'tech': 'Python, BERT, LSTM, Flask, SEBI API',
        'live': 'https://investguardai.onrender.com/',
        'github': 'https://github.com/JMadhan1/InvestGuard-AI'
    },
    {
        'name': '4. OneDeFi - Unified DeFi Management Hub',
        'badge': '🏆 IIIT Bangalore Winner | Live',
        'desc': 'Unified decentralized finance hub with cross-chain asset management and yield farming optimization',
        'tech': 'Python, Web3.py, Solana SDK, Flask, MCP Protocol',
        'live': 'https://onedefi-97wa.onrender.com/',
        'github': 'https://github.com/JMadhan1/ONEDeFi-MCP-Server'
    },
    {
        'name': '5. Navigate AI - Smart Navigation Platform',
        'badge': '🏆 Live',
        'desc': 'Smart AI navigation with intelligent travel planning and ML-powered route optimization',
        'tech': 'Python, Voice AI, NLP, Flask, ML Routing',
        'live': 'https://navigate-ai-2.onrender.com/',
        'github': 'https://github.com/JMadhan1/NAVIGATE-AI'
    },
    {
        'name': '6. CareerSync AI - Career Development Platform',
        'badge': '🏆 Alliance TechFest Winner | Live',
        'desc': 'AI-powered career platform with 85% automation, 70% faster hiring, 92% match accuracy',
        'tech': 'Python, NLP, TensorFlow, BERT, Flask',
        'live': 'https://careersync-ai.onrender.com/',
        'github': 'https://github.com/JMadhan1/AI-Powered-Resume-Job-matching-System'
    },
    {
        'name': '7. Smart Receipt Manager with Vision API',
        'badge': '🏆 Live',
        'desc': '95% OCR accuracy across 12 languages, 80% reduction in processing time, 99.99% uptime',
        'tech': 'Python, Flask, Google Vision API, Gemini AI, Firestore',
        'live': 'https://raseed-842527385575.us-central1.run.app/dashboard',
        'github': 'https://github.com/JMadhan1/Smart-Receipt-Manager'
    },
    {
        'name': '8. DocuMind - Abstractive Text Summarizer',
        'badge': '🏆 Live',
        'desc': 'ROUGE-L score of 0.85, 75% reduction in processing time, 8-second average processing',
        'tech': 'Python, BART Transformers, FastAPI, Streamlit, Azure AI',
        'live': 'https://documind-ai.onrender.com',
        'github': 'https://github.com/JMadhan1/DocuMind'
    },
    {
        'name': '9. Facial Authenticity Verification Model',
        'badge': '🏆 Research Project',
        'desc': '99.4% accuracy with SVM model, 98.5% with ResNet50, 81 unique image features',
        'tech': 'Python, TensorFlow, OpenCV, Scikit-learn, ResNet50',
        'live': None,
        'github': 'https://github.com/JMadhan1/Adaptive-Facial-identity-verification-system'
    },
    {
        'name': '10. AI-Powered TB Diagnosis Platform (PulmoScan.AI)',
        'badge': '🏆 Healthcare Innovation',
        'desc': '94% diagnostic accuracy, diagnosis time reduced from days to minutes',
        'tech': 'Python, TensorFlow Lite, Computer Vision, Healthcare AI',
        'live': None,
        'github': 'https://github.com/JMadhan1/Ai-Powered-TB-Daignosis-Care-Platform--pulmoscan.ai'
    }
]

for project in projects:
    # Project name
    p_name = doc.add_paragraph()
    p_name_run = p_name.add_run(project['name'])
    p_name_run.bold = True
    p_name_run.font.size = Pt(11)
    p_name_run.font.color.rgb = RGBColor(79, 70, 229)
    
    # Badge
    p_badge = doc.add_paragraph(project['badge'])
    p_badge.runs[0].font.size = Pt(9)
    p_badge.runs[0].italic = True
    
    # Description
    p_desc = doc.add_paragraph('• ' + project['desc'])
    p_desc.runs[0].font.size = Pt(9)
    
    # Tech Stack
    p_tech = doc.add_paragraph()
    p_tech_label = p_tech.add_run('Tech Stack: ')
    p_tech_label.bold = True
    p_tech_label.font.size = Pt(9)
    p_tech_text = p_tech.add_run(project['tech'])
    p_tech_text.font.size = Pt(9)
    
    # Links
    p_links = doc.add_paragraph()
    p_links_label = p_links.add_run('Links: ')
    p_links_label.bold = True
    p_links_label.font.size = Pt(9)
    
    links_text = []
    if project['live']:
        links_text.append(f"Live Demo: {project['live']}")
    if project['github']:
        links_text.append(f"GitHub: {project['github']}")
    
    p_links_text = p_links.add_run(' | '.join(links_text))
    p_links_text.font.size = Pt(8)
    p_links_text.font.color.rgb = RGBColor(37, 99, 235)
    
    doc.add_paragraph()  # Spacing

# Add page break
doc.add_page_break()

# Hackathon Achievements
doc.add_heading('HACKATHON ACHIEVEMENTS', 1)

hackathons = [
    {
        'title': '🥇 MIT Mysore Hackathon - 1st Place Winner (September 2025)',
        'details': [
            'Competition: 200+ teams from top 10 engineering colleges',
            'Project: CiteCraft.AI - AI-powered research paper generation platform',
            'Recognition: "Exceptional implementation of AI in academic research"'
        ]
    },
    {
        'title': '🥇 IIIT Bangalore AI Summit - Winner (August 2025)',
        'details': [
            'Project: ONEDeFi - Multi-Chain DeFi MCP Server',
            'Achievements: 23% Yield Increase, 99.9% Uptime, Multi-Chain Support'
        ]
    },
    {
        'title': '🥇 Alliance TechFest - 1st Place (March 2025)',
        'details': [
            'Competition: 150+ teams, Karnataka Regional',
            'Project: AI Resume & Job Matcher',
            'Impact: 85% Automation, 70% Faster Hiring, 92% Match Accuracy'
        ]
    },
    {
        'title': '🎓 IIT Bombay E-Yantra Fellowship (Aug 2024 - Mar 2025)',
        'details': [
            'Project: TutorTeach.ai - AI-powered adaptive learning platform',
            'Recognition: Official IIT Bombay Fellowship Certificate'
        ]
    }
]

for hackathon in hackathons:
    h_title = doc.add_paragraph()
    h_title_run = h_title.add_run(hackathon['title'])
    h_title_run.bold = True
    h_title_run.font.size = Pt(10)
    
    for detail in hackathon['details']:
        h_detail = doc.add_paragraph('  • ' + detail)
        h_detail.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()  # Spacing

# Professional Experience
doc.add_heading('PROFESSIONAL EXPERIENCE', 1)

experiences = [
    {
        'title': 'Summer Research Intern | MANIT Bhopal',
        'duration': 'Jun 2025 – Jul 2025',
        'points': [
            'Created AI-facilitated facial authenticity verification system with 99.4% accuracy',
            'Developed 81 unique image features and compared 7 ML and 4 DL architectures',
            'Achieved 98.5% accuracy with ResNet50 transfer learning model',
            'Tech: Python, TensorFlow, OpenCV, Scikit-learn, ResNet50'
        ]
    },
    {
        'title': 'AI Azure Intern | Microsoft (Edunet Foundation + AICTE)',
        'duration': 'May 2025 – Jun 2025',
        'points': [
            'Built DocuMind with ROUGE-L score of 0.85 and 8-second processing time',
            'Implemented abstractive summarization with multi-format support',
            'Recognized by AICTE and Edunet Foundation for successful completion',
            'Tech: Azure AI, Python, FastAPI, Streamlit, BART'
        ]
    },
    {
        'title': 'Fellowship Participant | IIT Bombay e-Yantra',
        'duration': 'Aug 2024 – Mar 2025',
        'points': [
            'Contributed to TutorTeach.ai - AI-powered adaptive learning platform',
            'Developed core features using Python and computer vision',
            'Earned IIT Bombay fellowship certification'
        ]
    }
]

for exp in experiences:
    e_title = doc.add_paragraph()
    e_title_run = e_title.add_run(exp['title'])
    e_title_run.bold = True
    e_title_run.font.size = Pt(11)
    
    e_duration = doc.add_paragraph(exp['duration'])
    e_duration.runs[0].font.size = Pt(9)
    e_duration.runs[0].italic = True
    
    for point in exp['points']:
        e_point = doc.add_paragraph('  • ' + point)
        e_point.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()  # Spacing

# Certifications
doc.add_heading('CERTIFICATIONS', 1)

certifications = [
    'AWS Certified Cloud Practitioner - Amazon Web Services (Dec 2024)',
    'Microsoft Certified: Azure AI Fundamentals - Microsoft (Nov 2024)',
    'Oracle AI Associate - Oracle (Oct 2024)',
    'Google AI Essentials - Google/Coursera (Sep 2024)',
    'NVIDIA AI Fundamentals - NVIDIA (Aug 2024)',
    'e-Yantra Entrepreneurship - IIT Bombay (Aug 2024 - Mar 2025)',
    'Forward Program - McKinsey & Company (Jul 2024)',
    'Internet of Things - NPTEL (IIT Madras) (Jun 2024)'
]

for cert in certifications:
    c = doc.add_paragraph('• ' + cert)
    c.runs[0].font.size = Pt(9)

# Vision & Interests
doc.add_heading('VISION & INTERESTS', 1)

visions = [
    ('🚀 AI-First Startups:', 'Building innovative AI companies solving real-world problems in healthcare, fintech, and education'),
    ('🔬 AI Research:', 'Advancing AI/ML innovation in computer vision, NLP, and generative AI with focus on ethical AI'),
    ('❤️ Healthcare AI:', 'Developing AI solutions that improve healthcare accessibility and diagnostic accuracy'),
    ('📚 EdTech Innovation:', 'Transforming education through AI-powered personalized learning')
]

for emoji_title, desc in visions:
    v = doc.add_paragraph()
    v_title = v.add_run(emoji_title)
    v_title.bold = True
    v_title.font.size = Pt(10)
    v_desc = v.add_run(' ' + desc)
    v_desc.font.size = Pt(9)

# Contact
doc.add_paragraph('_' * 100)
doc.add_heading('CONTACT', 1)

contact_info = [
    '📧 Email: jmadhan1@gmail.com',
    '🔗 GitHub: github.com/JMadhan1',
    '💼 LinkedIn: linkedin.com/in/jmadhan1',
    '🌐 Portfolio: jmadhan1.github.io/MadhanPortfolio'
]

for info in contact_info:
    c = doc.add_paragraph(info)
    c.runs[0].font.size = Pt(10)

# Footer quote
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('"Building the future of AI, one solution at a time"')
footer_run.italic = True
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(100, 116, 139)

# Save the document
doc.save('J_Madhan_Portfolio.docx')
print("Portfolio document created successfully: J_Madhan_Portfolio.docx")
